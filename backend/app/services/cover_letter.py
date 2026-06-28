"""
cover_letter.py
Generates a tailored cover letter using Claude, then saves as .docx.

The letter is personalised to:
  - The specific job and company (from jd_structured)
  - The candidate's actual projects and experience (from resume_data + matched_payload)
  - The selected keywords the user wanted to highlight
  - A tone the user can choose: Professional / Conversational / Concise

Output: .docx file (same naming convention as resume)
"""

import re
import json
import tempfile
from pathlib import Path

from app.llm import LLMError
from docx import Document
from docx.shared import Pt, Inches, Mm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement




def generate_cover_letter_text(
    jd_structured: dict,
    resume_data: dict,
    matched_payload: dict,
    selected_keywords: list,
    tone: str,
    llm,
    extra_instructions: str = "",
) -> str:
    """
    Ask Claude to write a tailored cover letter.
    Returns the letter as plain text (paragraphs separated by \\n\\n).
    """
    tone_guide = {
        "Professional": "Formal, polished, third-person distance. No contractions. Suit-and-tie energy.",
        "Conversational": "Warm, direct, first-person. Contractions OK. Human and approachable without being casual.",
        "Concise": "3 short paragraphs max. Every sentence earns its place. No filler phrases.",
    }.get(tone, "Professional, clear, and confident.")

    selected_projects = matched_payload.get("selected_projects", [])
    # Include project names, tech stack AND key bullet points for richer context
    proj_lines = []
    for p in selected_projects[:4]:
        proj_lines.append(f"- {p.get('name', '')} ({', '.join(p.get('tech_stack', [])[:5])})")
        for b in p.get("bullets", [])[:2]:  # top 2 bullets per project
            proj_lines.append(f"    • {b}")
    proj_summary = "\n".join(proj_lines)

    exp_summary = "\n".join(
        f"- {e.get('title', '')} @ {e.get('company', '')} ({e.get('dates', '')}): "
        + "; ".join(e.get("bullets", [])[:1])  # first bullet for context
        for e in resume_data.get("experience", [])[:3]
    )
    keyword_str = (
        ", ".join(selected_keywords[:12]) if selected_keywords else "none specified"
    )

    prompt = f"""You are an expert career coach and writer.

Write a tailored cover letter for this candidate applying to this role.

## JOB:
Role: {jd_structured.get("job_title", "")}
Company: {jd_structured.get("company", "")}
Location: {jd_structured.get("location", "")}
Key requirements: {", ".join(jd_structured.get("required_skills", [])[:10])}
Key responsibilities: {"; ".join(jd_structured.get("responsibilities", [])[:5])}

## CANDIDATE:
Name: {resume_data.get("name", "")}
Education: {", ".join(f"{e.get('degree', '')} from {e.get('school', '')}" for e in resume_data.get("education", [])[:2])}
Experience:
{exp_summary}

Resume projects (USE these specific details and metrics in the letter):
{proj_summary}

Keywords to emphasise: {keyword_str}

## TONE:
{tone_guide}

## STRICT RULES:
- Opening paragraph: Hook — why THIS company/role specifically, tied to the JD responsibilities
- Middle paragraph(s): 2-3 concrete examples using the EXACT project names and metrics above
- Closing paragraph: Clear call to action, confident not arrogant
- Do NOT use: "I am writing to express my interest", "I believe I am a strong candidate", "please find my resume attached", "thank you for your consideration", "I am excited to apply"
- Do NOT add a sign-off (Sincerely/Regards/Best) — it will be added automatically
- Do NOT include the candidate's name at the end
- Length: {"3 tight paragraphs" if tone == "Concise" else "4 paragraphs, ~350 words"}
- Address to: "Dear Hiring Team at {jd_structured.get("company", "[Company]")},"

{f"ADDITIONAL INSTRUCTIONS: {extra_instructions}" if extra_instructions else ""}

Return ONLY the cover letter body text — salutation through closing paragraph.
No subject line. No date. No sign-off. No name at the end."""

    try:
        return llm.complete(prompt=prompt, max_tokens=1200).text
    except LLMError as e:
        raise RuntimeError(f"Cover letter generation failed: {e}") from e


def _spacing(para, before=0, after=0):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    pPr.append(sp)


def _add_hyperlink_cl(para, text, url, size=9.0):
    """Insert a clickable hyperlink (reused from resume_builder pattern)."""
    from docx.opc.constants import RELATIONSHIP_TYPE as RT

    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)
    hl = OxmlElement("w:hyperlink")
    hl.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rSty = OxmlElement("w:rStyle")
    rSty.set(qn("w:val"), "Hyperlink")
    rPr.append(rSty)
    rFnt = OxmlElement("w:rFonts")
    rFnt.set(qn("w:ascii"), "Calibri")
    rFnt.set(qn("w:hAnsi"), "Calibri")
    rPr.append(rFnt)
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(int(size * 2)))
    rPr.append(sz_el)
    run.append(rPr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hl.append(run)
    para._p.append(hl)


def build_cover_letter_docx(
    letter_text: str,
    personal: dict,
    jd_structured: dict,
    output_dir: str = None,
    font_size: float = 10.5,
    bold_body: bool = False,
) -> dict:
    """
    Save the cover letter as a formatted .docx.
    Header matches resume exactly: large centred name, contact line with
    clickable LinkedIn + GitHub links, thin underline — then letter body.
    Same A4 page size and 0.5-inch margins as the resume.
    Returns {"docx_path": str, "pdf_path": str|None, "error": str|None}
    """
    result = {"docx_path": None, "pdf_path": None, "error": None}
    try:
        from docx.shared import Mm, Inches as _In

        doc = Document()

        for sec in doc.sections:
            sec.page_width = Mm(210)
            sec.page_height = Mm(297)
            sec.top_margin = _In(0.5)
            sec.bottom_margin = _In(0.5)
            sec.left_margin = _In(0.5)
            sec.right_margin = _In(0.5)

        BLACK = RGBColor(0x00, 0x00, 0x00)
        GREY = RGBColor(0x40, 0x40, 0x40)

        def _run(para, text, bold=False, italic=False, size=10.0, color=None):
            r = para.add_run(text)
            r.bold = bold
            r.italic = italic
            r.font.size = Pt(size)
            r.font.color.rgb = color or BLACK
            r.font.name = "Calibri"
            return r

        def _add_border_bottom(para, color="000000", sz=4):
            pPr = para._p.get_or_add_pPr()
            pBdr = OxmlElement("w:pBdr")
            bot = OxmlElement("w:bottom")
            bot.set(qn("w:val"), "single")
            bot.set(qn("w:sz"), str(sz))
            bot.set(qn("w:space"), "1")
            bot.set(qn("w:color"), color)
            pBdr.append(bot)
            pPr.append(pBdr)

        p_name = doc.add_paragraph()
        p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p_name, before=0, after=10)
        r = p_name.add_run(personal.get("name", ""))
        r.bold = True
        r.font.size = Pt(20)
        r.font.color.rgb = BLACK
        r.font.name = "Calibri"

        p_contact = doc.add_paragraph()
        p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _spacing(p_contact, before=0, after=60)
        _add_border_bottom(p_contact, color="000000", sz=4)

        cs = 9.0
        sep = "  |  "
        parts = []
        if personal.get("location"):
            parts.append(personal["location"])
        if personal.get("phone"):
            parts.append(personal["phone"])
        if personal.get("email"):
            parts.append(personal["email"])

        if parts:
            _run(p_contact, sep.join(parts), size=cs, color=GREY)

        li_url = personal.get("linkedin_url", "")
        if not li_url and personal.get("linkedin", ""):
            raw = personal["linkedin"]
            if "linkedin.com" in raw:
                li_url = ("https://" + raw) if not raw.startswith("http") else raw
            else:
                li_url = "https://www.linkedin.com/in/" + raw if raw else ""
        if li_url:
            _run(p_contact, sep + li_url, size=cs, color=GREY)

        gh_url = personal.get("github_url", "")
        if not gh_url and personal.get("github", ""):
            raw = personal["github"]
            if "github.com" in raw:
                gh_url = ("https://" + raw) if not raw.startswith("http") else raw
            else:
                gh_url = "https://github.com/" + raw if raw else ""
        if gh_url:
            _run(p_contact, sep + gh_url, size=cs, color=GREY)

        SIGNOFF_PATTERNS = [
            "sincerely,",
            "best regards,",
            "regards,",
            "respectfully,",
            "best,",
            "yours truly,",
        ]
        letter_clean = letter_text.strip()
        # Remove trailing sign-off + name if Claude included it
        lines = letter_clean.split("\n")
        # Walk backwards and drop lines that match sign-off patterns
        while lines and any(
            lines[-1].strip().lower().startswith(s) for s in SIGNOFF_PATTERNS
        ):
            lines.pop()
        # Also drop any trailing line that exactly matches the candidate's name
        cand_name = personal.get("name", "").strip().lower()
        if lines and lines[-1].strip().lower() == cand_name:
            lines.pop()
        letter_clean = "\n".join(lines).strip()

        paragraphs = [p.strip() for p in letter_clean.split("\n\n") if p.strip()]
        for i, para_text in enumerate(paragraphs):
            p = doc.add_paragraph()
            _spacing(p, before=0, after=120 if i < len(paragraphs) - 1 else 200)
            # Justify body text
            pPr = p._p.get_or_add_pPr()
            jc = OxmlElement("w:jc")
            jc.set(qn("w:val"), "both")
            pPr.append(jc)
            r2 = p.add_run(para_text.replace("\n", " "))
            r2.font.size = Pt(font_size)
            r2.bold = bold_body
            r2.font.color.rgb = BLACK
            r2.font.name = "Calibri"

        p_sig = doc.add_paragraph()
        _spacing(p_sig, before=200, after=0)
        _run(p_sig, "Sincerely,", size=font_size)
        p_name2 = doc.add_paragraph()
        _spacing(p_name2, before=40, after=0)
        _run(p_name2, personal.get("name", ""), bold=True, size=font_size)

        # Filename
        def _clean(s):
            s = re.sub(r"[^\w\s\-]", "", s or "").strip()
            return re.sub(r"[\s\-]+", "_", s)

        name_parts = (personal.get("name", "CoverLetter") or "CoverLetter").split()
        lastname = _clean(name_parts[-1]) if len(name_parts) > 1 else ""
        firstname = _clean(name_parts[0]) if name_parts else "Cover"
        company = _clean(jd_structured.get("company", ""))
        fn = "_".join(p for p in [lastname, firstname, company, "Cover_Letter"] if p)
        fn = re.sub(r"_+", "_", fn).strip("_")

        out = (
            Path(output_dir)
            if output_dir
            else Path(tempfile.mkdtemp(prefix="resumeforge_cl_"))
        )
        out.mkdir(parents=True, exist_ok=True)
        path = out / f"{fn}.docx"
        doc.save(str(path))
        result["docx_path"] = str(path)

        # PDF conversion — run in subprocess to avoid COM conflicts with resume builder
        pdf_path = out / f"{fn}.pdf"
        try:
            import time, subprocess, sys as _sys

            time.sleep(1.0)
            proc = subprocess.run(
                [
                    _sys.executable,
                    "-c",
                    f"from docx2pdf import convert; convert(r'{path}', r'{pdf_path}')",
                ],
                capture_output=True,
                text=True,
                timeout=60,
            )
            if proc.returncode == 0 and pdf_path.exists():
                result["pdf_path"] = str(pdf_path)
            else:
                from docx2pdf import convert

                convert(str(path), str(pdf_path))
                if pdf_path.exists():
                    result["pdf_path"] = str(pdf_path)
                else:
                    result["pdf_error"] = proc.stderr[:200]
        except Exception as e:
            result["pdf_error"] = str(e)

    except Exception as e:
        result["error"] = str(e)

    return result


def revise_cover_letter(
    letter_text: str,
    edit_instructions: str,
    jd_structured: dict,
    resume_data: dict,
    llm,
) -> str:
    """Apply edit instructions to an existing cover letter text."""
    prompt = f"""You are a professional cover letter editor.
Apply ONLY the requested edits to this cover letter. Do not change anything not mentioned.

CURRENT COVER LETTER:
{letter_text}

EDIT INSTRUCTIONS:
{edit_instructions}

Return the complete updated cover letter text. No explanation, no markdown. Just the letter."""

    try:
        return llm.complete(prompt=prompt, max_tokens=1200).text
    except LLMError as e:
        raise RuntimeError(f"Cover letter revision failed: {e}") from e
