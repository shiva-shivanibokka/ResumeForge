"""
resume_parser.py
Extracts structured data from an uploaded resume (PDF or .docx).

Returns a dict with:
  name, email, phone, linkedin, github, location,
  education (list), experience (list), skills (list), raw_text
"""

import io
import re
from pathlib import Path

# TEXT EXTRACTION


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF using PyMuPDF (fast, layout-aware)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        return "\n".join(pages)
    except Exception as e:
        # Fallback to pdfminer
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams

            output = io.StringIO()
            extract_text_to_fp(
                io.BytesIO(file_bytes),
                output,
                laparams=LAParams(),
                output_type="text",
                codec="utf-8",
            )
            return output.getvalue()
        except Exception as e2:
            return f"[PDF extraction failed: {e} / {e2}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """
    Extract plain text from a .docx file.
    Also resolves hyperlinks — replaces display text like 'LinkedIn' or 'GitHub'
    with the actual underlying URL so the parser can extract the real link.
    """
    try:
        import docx
        from docx.oxml.ns import qn

        doc = docx.Document(io.BytesIO(file_bytes))

        def _para_text_with_hyperlinks(para) -> str:
            """
            Return paragraph text, substituting hyperlink display text with the
            actual URL wherever the underlying URL is a LinkedIn or GitHub link.
            """
            # Relationships for this paragraph's part
            rels = para.part.rels

            result = []
            for child in para._p:
                tag = child.tag.split("}")[-1] if "}" in child.tag else child.tag
                if tag == "hyperlink":
                    # Get the relationship id
                    r_id = child.get(qn("r:id"))
                    url = ""
                    if r_id and r_id in rels:
                        url = rels[r_id].target_ref or ""
                    # Get display text
                    display = "".join(t.text for t in child.iter(qn("w:t")) if t.text)
                    # If the URL is a LinkedIn or GitHub link, emit the URL instead
                    if url and ("linkedin.com" in url or "github.com" in url):
                        result.append(url)
                    else:
                        result.append(display)
                elif tag == "r":
                    # Normal run — grab text
                    for t in child.iter(qn("w:t")):
                        if t.text:
                            result.append(t.text)
            return "".join(result)

        paragraphs = []
        for para in doc.paragraphs:
            text = _para_text_with_hyperlinks(para).strip()
            if text:
                paragraphs.append(text)

        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    cell_text = ""
                    for para in cell.paragraphs:
                        cell_text += _para_text_with_hyperlinks(para)
                    if cell_text.strip():
                        paragraphs.append(cell_text.strip())

        return "\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def extract_raw_text(file_path: str, file_bytes: bytes = None) -> str:
    """
    Extract raw text from a resume file.
    Accepts either a file path string or raw bytes + extension hint.
    """
    path = Path(file_path)
    if file_bytes is None:
        file_bytes = path.read_bytes()

    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        # Try to decode as plain text
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


# REGEX FIELD EXTRACTION  (fast, no API call needed)


def _find_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def _find_phone(text: str) -> str:
    match = re.search(r"(\+?1[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}", text)
    return match.group(0).strip() if match else ""


def _find_linkedin(text: str) -> str:
    match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    return match.group(0) if match else ""


def _find_github(text: str) -> str:
    match = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    return match.group(0) if match else ""


def _find_name(text: str) -> str:
    """
    Heuristic: the name is usually on the first non-empty line,
    often as 2-4 capitalized words.
    """
    for line in text.splitlines()[:10]:
        line = line.strip()
        if not line:
            continue
        # Skip lines that look like contact info
        if any(c in line for c in ["@", "http", "linkedin", "github", "+1", "("]):
            continue
        # Match 2-4 capitalized words (name pattern)
        if re.match(r"^([A-Z][a-z]+\s){1,3}[A-Z][a-z]+$", line):
            return line
        # Also accept ALL-CAPS names
        if re.match(r"^[A-Z\s]{4,40}$", line) and len(line.split()) >= 2:
            return line.title()
    return ""


# CLAUDE-POWERED STRUCTURED EXTRACTION


def parse_resume_with_claude(raw_text: str, llm) -> dict:
    """
    Use Claude to extract fully structured data from resume text.
    Falls back to regex for contact info if Claude fails.
    """
    import json

    prompt = f"""You are a resume parsing assistant. Extract structured information from this resume text.

Return a JSON object with exactly these keys:
{{
  "name": "Full name",
  "email": "email address",
  "phone": "phone number",
  "linkedin": "LinkedIn URL or username",
  "github": "GitHub URL or username",
  "location": "City, State or City, Country",
  "education": [
    {{
      "school": "Institution name",
      "degree": "Degree and field of study",
      "dates": "Start – End",
      "gpa": "GPA if listed, else null",
      "location": "City, State"
    }}
  ],
  "experience": [
    {{
      "company": "Company name",
      "title": "Job title",
      "dates": "Start – End",
      "location": "City, State",
      "bullets": ["bullet 1", "bullet 2", ...]
    }}
  ],
  "skills": ["skill1", "skill2", ...],
  "summary": "Professional summary if present, else empty string"
}}

Only return the JSON. No markdown fences. No explanation.

RESUME TEXT:
{raw_text[:8000]}"""

    try:
        raw = llm.complete(prompt=prompt, max_tokens=3000).text
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
        parsed = json.loads(raw)
    except Exception:
        # Fallback: build minimal structure from regex
        parsed = {
            "name": _find_name(raw_text),
            "email": _find_email(raw_text),
            "phone": _find_phone(raw_text),
            "linkedin": _find_linkedin(raw_text),
            "github": _find_github(raw_text),
            "location": "",
            "education": [],
            "experience": [],
            "skills": [],
            "summary": "",
        }

    # Always patch contact fields with regex if Claude missed them or returned
    # display text ("LinkedIn", "GitHub") instead of the actual URL
    if not parsed.get("email"):
        parsed["email"] = _find_email(raw_text)
    if not parsed.get("phone"):
        parsed["phone"] = _find_phone(raw_text)
    # Prefer regex URL over Claude's value if it looks like a real URL
    regex_li = _find_linkedin(raw_text)
    if regex_li and "linkedin.com" in regex_li:
        parsed["linkedin"] = regex_li
    elif not parsed.get("linkedin"):
        parsed["linkedin"] = regex_li
    regex_gh = _find_github(raw_text)
    if regex_gh and "github.com" in regex_gh:
        parsed["github"] = regex_gh
    elif not parsed.get("github"):
        parsed["github"] = regex_gh
    if not parsed.get("name"):
        parsed["name"] = _find_name(raw_text)

    parsed["raw_text"] = raw_text
    return parsed


def parse_resume(file_path: str, llm, file_bytes: bytes = None) -> dict:
    """
    Full pipeline: extract text → parse with Claude → return structured dict.
    """
    raw_text = extract_raw_text(file_path, file_bytes)
    # Match the specific failure sentinels, not any leading "[" — a résumé whose
    # text legitimately starts with "[" must not be misread as an extraction failure.
    _FAIL = ("[PDF extraction failed:", "[DOCX extraction failed:")
    if not raw_text or raw_text.startswith(_FAIL):
        return {
            "name": "",
            "email": "",
            "phone": "",
            "linkedin": "",
            "github": "",
            "location": "",
            "education": [],
            "experience": [],
            "skills": [],
            "summary": "",
            "raw_text": raw_text,
            "error": raw_text,
        }
    return parse_resume_with_claude(raw_text, llm)
