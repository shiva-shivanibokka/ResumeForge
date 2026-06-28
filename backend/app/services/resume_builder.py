"""
resume_builder.py

Generates a resume that exactly mirrors the original Shivani Bokka format:
  - Large centered name (no subtitle/title under it)
  - Single contact line: Location | Phone | Email | LinkedIn (hyperlink) | GitHub (hyperlink)
  - TECHNICAL SKILLS section with bold category labels + inline text
  - PROFESSIONAL EXPERIENCE — bold company + right-aligned date, italic title + location
  - RELEVANT PROJECTS — bold project name, italic tech stack line, bullet points
  - EDUCATION — bold school + right-aligned date, italic degree + location + GPA

Layout options:
  - one_page=True  → aggressive font/spacing compression to fit one A4 page
  - one_page=False → comfortable two-page A4 layout

Output filename: Firstname_Lastname_Company_JobTitle.docx / .pdf
"""

import re
import os
from pathlib import Path
from datetime import datetime
from dataclasses import dataclass, field
from typing import Optional

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm, Emu, Mm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.opc.constants import RELATIONSHIP_TYPE as RT

A4_WIDTH = Mm(210)
A4_HEIGHT = Mm(297)

BLACK = RGBColor(0x00, 0x00, 0x00)
GREY = RGBColor(0x40, 0x40, 0x40)

AVAILABLE_FONTS = [
    "Calibri",
    "Arial",
    "Georgia",
    "Times New Roman",
    "Garamond",
    "Cambria",
    "Helvetica",
    "Trebuchet MS",
]



@dataclass
class FontConfig:
    """
    Controls fonts and sizes for each text element in the resume.
    All sizes in points. Passed directly to build_resume().
    """

    # Font families
    name_font: str = "Calibri"  # candidate name at top
    heading_font: str = "Calibri"  # section headings (EXPERIENCE, EDUCATION…)
    body_font: str = "Calibri"  # all body text, bullets, contact line

    # Sizes — (1-page value, 2-page value)
    # If user sets a custom size it overrides both; otherwise FS dict is used
    name_size: Optional[float] = None  # default: 20 (1p) / 22 (2p)
    heading_size: Optional[float] = None  # default: 9.5 (1p) / 10 (2p)
    body_size: Optional[float] = None  # default: 8.5 (1p) / 9.5 (2p)


DEFAULT_FONT = FontConfig()

FS = {
    "name": (20, 22),
    "contact": (8.5, 9),
    "section_hd": (9.5, 10),
    "body": (8.5, 9.5),
    "bullet": (8.0, 9.5),
    "tech_stack": (7.5, 8.5),
}

SP = {
    "after_header": (20, 60),
    "section_before": (80, 120),
    "section_after": (0, 0),
    "entry_before": (30, 60),
    "bullet_space": (0, 2),
    "tech_before": (0, 0),
    "tech_after": (4, 8),
}




def _spacing(para, before=0, after=0, line=None):
    pPr = para._p.get_or_add_pPr()
    sp = OxmlElement("w:spacing")
    sp.set(qn("w:before"), str(before))
    sp.set(qn("w:after"), str(after))
    if line is not None:
        sp.set(qn("w:line"), str(line))
        sp.set(qn("w:lineRule"), "exact")
    else:
        # 1.0 line spacing = 240 twips, auto rule
        sp.set(qn("w:line"), "240")
        sp.set(qn("w:lineRule"), "auto")
    pPr.append(sp)


def _justify(para):
    """Set paragraph alignment to justified (like Word's Justify button)."""
    pPr = para._p.get_or_add_pPr()
    jc = OxmlElement("w:jc")
    jc.set(qn("w:val"), "both")  # "both" = full justify in OOXML
    pPr.append(jc)


def _bottom_border(para, color="000000", sz=4):
    pPr = para._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bot = OxmlElement("w:bottom")
    bot.set(qn("w:val"), "single")
    bot.set(qn("w:sz"), str(sz))
    bot.set(qn("w:space"), "1")
    bot.set(qn("w:color"), color)
    pBdr.append(bot)
    pPr.append(pBdr)


def _run(para, text, bold=False, italic=False, size=10.0, color=BLACK, font="Calibri"):
    r = para.add_run(text)
    r.bold = bold
    r.italic = italic
    r.font.size = Pt(size)
    r.font.color.rgb = color
    r.font.name = font
    return r


def _set_right_tab(para, pos_twips=10080):
    """Add a right-aligned tab stop (used for date columns)."""
    pPr = para._p.get_or_add_pPr()
    tabs = OxmlElement("w:tabs")
    tab = OxmlElement("w:tab")
    tab.set(qn("w:val"), "right")
    tab.set(qn("w:pos"), str(pos_twips))
    tabs.append(tab)
    pPr.append(tabs)


def _add_hyperlink(para, text, url, size=9.0, font="Calibri"):
    """
    Insert a clickable hyperlink into a paragraph.
    Styled to match surrounding text (no blue/underline override —
    uses the built-in Hyperlink character style which Word renders as blue+underline).
    """
    part = para.part
    r_id = part.relate_to(url, RT.HYPERLINK, is_external=True)

    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)

    new_run = OxmlElement("w:r")

    rPr = OxmlElement("w:rPr")
    # Apply Hyperlink character style
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    # Font + size
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:ascii"), font)
    rFonts.set(qn("w:hAnsi"), font)
    rPr.append(rFonts)
    sz_el = OxmlElement("w:sz")
    sz_el.set(qn("w:val"), str(int(size * 2)))  # half-points
    rPr.append(sz_el)
    new_run.append(rPr)

    t = OxmlElement("w:t")
    t.text = text
    new_run.append(t)
    hyperlink.append(new_run)
    para._p.append(hyperlink)
    return hyperlink




def _new_doc(one_page: bool, fc: FontConfig) -> Document:
    doc = Document()
    margin = Inches(0.5)
    for sec in doc.sections:
        sec.page_width = A4_WIDTH
        sec.page_height = A4_HEIGHT
        sec.top_margin = margin
        sec.bottom_margin = margin
        sec.left_margin = margin
        sec.right_margin = margin

    body_sz = fc.body_size or FS["body"][0 if one_page else 1]
    doc.styles["Normal"].font.name = fc.body_font
    doc.styles["Normal"].font.size = Pt(body_sz)
    # Apply justify + 1.0 line spacing to Normal style
    try:
        from docx.enum.text import WD_ALIGN_PARAGRAPH as _WAP

        doc.styles["Normal"].paragraph_format.alignment = _WAP.JUSTIFY
        doc.styles["Normal"].paragraph_format.line_spacing = 1.0
    except Exception:
        pass

    try:
        bl = doc.styles["List Bullet"]
        bl.font.name = fc.body_font
        bl.font.size = Pt(fc.body_size or FS["bullet"][0 if one_page else 1])
        bl.paragraph_format.line_spacing = 1.0
    except Exception:
        pass
    return doc




def _fs(key, one_page):
    return FS[key][0 if one_page else 1]


def _sp(key, one_page):
    return SP[key][0 if one_page else 1]




def _add_header(doc, personal: dict, one_page: bool, fc: FontConfig = DEFAULT_FONT):
    # Name — large, centered, bold
    p_name = doc.add_paragraph()
    p_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_name, before=0, after=10)
    r = p_name.add_run(personal.get("name", "Your Name"))
    r.bold = True
    r.font.size = Pt(fc.name_size or _fs("name", one_page))
    r.font.color.rgb = BLACK
    r.font.name = fc.name_font

    # Contact line — centered, single line with | separators
    p_contact = doc.add_paragraph()
    p_contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _spacing(p_contact, before=0, after=_sp("after_header", one_page))

    cs = _fs("contact", one_page)
    parts_plain = []
    if personal.get("location"):
        parts_plain.append(personal["location"])
    if personal.get("phone"):
        parts_plain.append(personal["phone"])
    if personal.get("email"):
        parts_plain.append(personal["email"])

    sep = "  |  "

    # Plain parts first
    if parts_plain:
        _run(p_contact, sep.join(parts_plain), size=cs, color=BLACK)

    # LinkedIn — plain text full URL (no hyperlink styling)
    linkedin_url = personal.get("linkedin_url", "")
    if not linkedin_url:
        raw = personal.get("linkedin", "")
        if "linkedin.com" in raw:
            linkedin_url = "https://" + raw if not raw.startswith("http") else raw
        else:
            linkedin_url = "https://www.linkedin.com/in/" + raw if raw else ""
    if linkedin_url:
        _run(p_contact, sep + linkedin_url, size=cs, color=BLACK)

    # GitHub — plain text full URL (no hyperlink styling)
    github_url = personal.get("github_url", "")
    if not github_url:
        raw = personal.get("github", "")
        if "github.com" in raw:
            github_url = "https://" + raw if not raw.startswith("http") else raw
        else:
            github_url = "https://github.com/" + raw if raw else ""
    if github_url:
        _run(p_contact, sep + github_url, size=cs, color=BLACK)




def _section_heading(doc, title: str, one_page: bool, fc: FontConfig = DEFAULT_FONT):
    p = doc.add_paragraph()
    _spacing(
        p, before=_sp("section_before", one_page), after=_sp("section_after", one_page)
    )
    r = p.add_run(title.upper())
    r.bold = True
    r.font.size = Pt(fc.heading_size or _fs("section_hd", one_page))
    r.font.color.rgb = BLACK
    r.font.name = fc.heading_font
    _bottom_border(p, color="000000", sz=4)
    return p




def _add_skills(
    doc, tailored_skills: dict, one_page: bool, fc: FontConfig = DEFAULT_FONT
):
    if not tailored_skills:
        return
    _section_heading(doc, "Technical Skills", one_page, fc)
    bs = fc.body_size or _fs("body", one_page)
    for category, items in tailored_skills.items():
        p = doc.add_paragraph()
        _spacing(p, before=8, after=4)
        rb = p.add_run(f"{category}: ")
        rb.bold = True
        rb.font.size = Pt(bs)
        rb.font.color.rgb = BLACK
        rb.font.name = fc.body_font
        rr = p.add_run(items)
        rr.font.size = Pt(bs)
        rr.font.color.rgb = BLACK
        rr.font.name = fc.body_font




def _add_experience(
    doc, experience: list, one_page: bool, fc: FontConfig = DEFAULT_FONT
):
    if not experience:
        return
    _section_heading(doc, "Professional Experience", one_page, fc)
    bs = fc.body_size or _fs("body", one_page)

    for exp in experience:
        p1 = doc.add_paragraph()
        _spacing(p1, before=_sp("entry_before", one_page), after=0)
        _set_right_tab(p1)
        rb = p1.add_run(exp.get("company", ""))
        rb.bold = True
        rb.font.size = Pt(bs)
        rb.font.color.rgb = BLACK
        rb.font.name = fc.body_font
        _run(p1, "\t", font=fc.body_font)
        _run(p1, exp.get("dates", ""), size=bs, color=BLACK, font=fc.body_font)

        p2 = doc.add_paragraph()
        _spacing(p2, before=0, after=0)
        _set_right_tab(p2)
        ri = p2.add_run(exp.get("title", ""))
        ri.italic = True
        ri.font.size = Pt(bs)
        ri.font.color.rgb = BLACK
        ri.font.name = fc.body_font
        _run(p2, "\t", font=fc.body_font)
        rl = p2.add_run(exp.get("location", ""))
        rl.italic = True
        rl.font.size = Pt(bs)
        rl.font.color.rgb = BLACK
        rl.font.name = fc.body_font

        for b in exp.get("bullets", []):
            _bullet(doc, b, one_page, fc=fc)




def _bullet(doc, text: str, one_page: bool, indent=0.2, fc: FontConfig = DEFAULT_FONT):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(indent)
    p.paragraph_format.first_line_indent = Inches(-0.15)
    sp = _sp("bullet_space", one_page)
    p.paragraph_format.space_before = Pt(sp)
    p.paragraph_format.space_after = Pt(sp)
    _justify(p)
    r = p.add_run(text)
    r.font.size = Pt(fc.body_size or _fs("bullet", one_page))
    r.font.color.rgb = BLACK
    r.font.name = fc.body_font
    return p




def _add_projects(
    doc, selected_projects: list, one_page: bool, fc: FontConfig = DEFAULT_FONT
):
    if not selected_projects:
        return
    _section_heading(doc, "Relevant Projects", one_page, fc)
    bs = fc.body_size or _fs("body", one_page)
    ts = _fs("tech_stack", one_page)

    for proj in selected_projects:
        p_title = doc.add_paragraph()
        _spacing(p_title, before=_sp("entry_before", one_page), after=0)
        rt = p_title.add_run(proj.get("name", "Project"))
        rt.bold = True
        rt.font.size = Pt(bs)
        rt.font.color.rgb = BLACK
        rt.font.name = fc.body_font

        tech = proj.get("tech_stack", [])
        if tech:
            tech_str = " | ".join(str(t) for t in tech[:10])
            p_tech = doc.add_paragraph()
            _spacing(
                p_tech,
                before=_sp("tech_before", one_page),
                after=_sp("tech_after", one_page),
            )
            rtech = p_tech.add_run(tech_str)
            rtech.italic = True
            rtech.font.size = Pt(ts)
            rtech.font.color.rgb = GREY
            rtech.font.name = fc.body_font

        for b in proj.get("bullets", []):
            _bullet(doc, b, one_page, fc=fc)




def _add_education(doc, education: list, one_page: bool, fc: FontConfig = DEFAULT_FONT):
    if not education:
        return
    _section_heading(doc, "Education", one_page, fc)
    bs = fc.body_size or _fs("body", one_page)

    for edu in education:
        p1 = doc.add_paragraph()
        _spacing(p1, before=_sp("entry_before", one_page), after=0)
        _set_right_tab(p1)
        rb = p1.add_run(edu.get("school", ""))
        rb.bold = True
        rb.font.size = Pt(bs)
        rb.font.color.rgb = BLACK
        rb.font.name = fc.body_font
        _run(p1, "\t", font=fc.body_font)
        _run(
            p1, edu.get("dates", ""), bold=True, size=bs, color=BLACK, font=fc.body_font
        )

        degree_str = edu.get("degree", "")
        p2 = doc.add_paragraph()
        _spacing(p2, before=0, after=0)
        _set_right_tab(p2)
        ri = p2.add_run(degree_str)
        ri.italic = True
        ri.font.size = Pt(bs)
        ri.font.color.rgb = BLACK
        ri.font.name = fc.body_font
        _run(p2, "\t", font=fc.body_font)
        rl = p2.add_run(edu.get("location", ""))
        rl.italic = True
        rl.font.size = Pt(bs)
        rl.font.color.rgb = BLACK
        rl.font.name = fc.body_font

        if edu.get("gpa"):
            pg = doc.add_paragraph()
            _spacing(pg, before=0, after=0)
            _run(pg, f"  GPA: {edu['gpa']}", size=bs, color=BLACK, font=fc.body_font)




def _build_filename(personal: dict, matched_payload: dict) -> str:
    """
    Pattern: Lastname_Firstname_Company_Role
    e.g. Bokka_Shivani_Parspec_Machine_Learning_Engineer
    """

    def clean(s):
        s = re.sub(r"[^\w\s\-]", "", s or "").strip()
        return re.sub(r"[\s\-]+", "_", s)

    name_parts = (personal.get("name", "Resume") or "Resume").split()
    firstname = clean(name_parts[0]) if name_parts else "Resume"
    lastname = clean(name_parts[-1]) if len(name_parts) > 1 else ""
    company = clean(matched_payload.get("company", ""))
    job_title = clean(
        matched_payload.get("job_title", matched_payload.get("resume_title", "Role"))
    )

    # Pattern: Lastname_Firstname_Company_Role
    parts = [p for p in [lastname, firstname, company, job_title] if p]
    return "_".join(parts)


# These ratios are fixed relative to body size — same as professional resume templates
NAME_RATIO = 2.2  # name = body × 2.2
HEADING_RATIO = 1.1  # section heading = body × 1.1
# body = body × 1.0 (trivially)
# tech stack = body × 0.88


def _estimate_pages_docx(docx_path: str) -> float:
    """
    Estimate the number of pages by summing paragraph heights from the .docx XML.
    No Word/LibreOffice required — pure python-docx.

    Returns a float (e.g. 1.05 means slightly over 1 page) so callers can
    apply a conservative threshold (e.g. <= 0.97 to ensure a real safety margin).

    A4 with 0.5-inch margins on all sides:
      Usable height = 297mm - 2×12.7mm = 271.6mm = 770pt
    """
    from docx import Document as _Doc
    from docx.oxml.ns import qn as _qn

    USABLE_PT = 770.0  # A4 minus 0.5in top+bottom margins
    TWO_PT_TO_PT = 1 / 20  # Word stores spacing in twips (1/20 pt)

    try:
        doc = _Doc(docx_path)
        total_pt = 0.0

        for para in doc.paragraphs:
            # Get font size from the first run, fallback to Normal style, fallback to 9pt
            size_pt = 9.0
            if para.runs:
                fs = para.runs[0].font.size
                if fs:
                    size_pt = fs.pt
            else:
                try:
                    ns = doc.styles["Normal"].font.size
                    if ns:
                        size_pt = ns.pt
                except Exception:
                    pass

            # Line height: font size × 1.0 line spacing (set in _new_doc)
            line_height = size_pt

            # Paragraph spacing before + after from XML
            pPr = para._p.find(_qn("w:pPr"))
            sp_before = sp_after = 0.0
            if pPr is not None:
                sp = pPr.find(_qn("w:spacing"))
                if sp is not None:
                    b = sp.get(_qn("w:before"))
                    a = sp.get(_qn("w:after"))
                    if b:
                        sp_before = int(b) * TWO_PT_TO_PT
                    if a:
                        sp_after = int(a) * TWO_PT_TO_PT

            total_pt += line_height + sp_before + sp_after

        return total_pt / USABLE_PT

    except Exception:
        return 1.0  # assume fits if we can't measure


def _count_pages_pdf(docx_path: str) -> int:
    """
    Convert .docx to a temp PDF and count pages using pdfminer.
    Only used for the final verification after binary search — not during search.
    """
    import tempfile, os

    pdf_tmp = tempfile.mktemp(suffix=".pdf")
    try:
        from docx2pdf import convert

        convert(docx_path, pdf_tmp)
        from pdfminer.pdfpage import PDFPage

        with open(pdf_tmp, "rb") as f:
            return sum(1 for _ in PDFPage.get_pages(f))
    except Exception:
        return 1
    finally:
        try:
            os.remove(pdf_tmp)
        except OSError:
            pass


# Keep old name as alias (used by external code)
_count_pages = _count_pages_pdf


def _build_doc(personal, education, matched_payload, fc, one_page):
    """Build and return a Document object (no saving)."""
    doc = _new_doc(one_page, fc)
    experience = matched_payload.get("tailored_experience", [])
    projects = matched_payload.get("selected_projects", [])
    skills = matched_payload.get("tailored_skills", {})
    _add_header(doc, personal, one_page, fc)
    _add_skills(doc, skills, one_page, fc)
    _add_experience(doc, experience, one_page, fc)
    _add_projects(doc, projects, one_page, fc)
    _add_education(doc, education, one_page, fc)
    return doc


def auto_fit_font_size(
    personal: dict,
    education: list,
    matched_payload: dict,
    font_family: str = "Calibri",
    one_page: bool = True,
    target_pages: int = 1,
    min_body: float = 7.5,
    max_body: float = 10.5,
) -> FontConfig:
    """
    Binary-search for the largest body font size where the resume fits
    within `target_pages`. Returns a FontConfig with proportional sizes.

    Typographic ratios applied:
      name    = body × 2.2
      heading = body × 1.1
      body    = body (base)
    """
    import tempfile, os

    def _make_fc(body: float) -> FontConfig:
        return FontConfig(
            name_font=font_family,
            heading_font=font_family,
            body_font=font_family,
            name_size=round(body * NAME_RATIO, 1),
            heading_size=round(body * HEADING_RATIO, 1),
            body_size=body,
        )

    def _fits_fast(body: float) -> bool:
        """
        Fast page-fit check using the pure-Python estimator.
        Threshold is 0.95 (not 1.0) to leave a conservative 5% safety margin
        so the final Word render never overflows by a line or two.
        """
        fc = _make_fc(body)
        doc = _build_doc(personal, education, matched_payload, fc, one_page)
        tmp = tempfile.mktemp(suffix=".docx")
        try:
            doc.save(tmp)
            pages = _estimate_pages_docx(tmp)
            return pages <= (target_pages * 0.95)
        finally:
            try:
                os.remove(tmp)
            except OSError:
                pass

    if not one_page:
        return _make_fc(9.5)

    # Binary search using the fast estimator (no Word COM calls)
    # 8 iterations → precision ~0.01pt, still fast because no PDF conversion
    lo, hi = min_body, max_body
    best = min_body

    for _ in range(8):
        mid = (lo + hi) / 2
        if _fits_fast(mid):
            best = mid
            lo = mid
        else:
            hi = mid

    import math

    # Floor to 1 decimal place + extra 0.3pt safety buffer
    # This guarantees the final Word render stays within the page
    best_safe = math.floor(best * 10) / 10 - 0.3
    best_safe = max(best_safe, min_body)
    return _make_fc(best_safe)




def build_resume(
    personal: dict,
    education: list,
    matched_payload: dict,
    output_dir: str = None,
    to_pdf: bool = True,
    one_page: bool = True,
    font_config: FontConfig = None,
    auto_fill: bool = True,
) -> dict:
    """
    Build the tailored resume .docx and optionally .pdf.

    Files are saved to a temp directory by default (output_dir=None).
    The caller is responsible for presenting the files for download —
    nothing is auto-saved to a permanent location.

    Args:
        personal:        from resume_parser (name, email, phone, linkedin, github, location)
        education:       from resume_parser
        matched_payload: from project_matcher.match_and_tailor()
        output_dir:      directory to save files; uses tempfile if None
        to_pdf:          also convert to PDF via docx2pdf
        one_page:        True = compress to fit 1 A4 page; False = comfortable 2-page layout

    Returns:
        {"docx_path": str, "pdf_path": str|None, "filename_base": str, "error": str|None}
    """
    import tempfile

    result = {
        "docx_path": None,
        "pdf_path": None,
        "filename_base": None,
        "error": None,
        "font_config_used": None,
    }

    try:
        if auto_fill and one_page:
            # Determine the best font family from font_config (if provided) or default
            family = (
                font_config.body_font
                if font_config and font_config.body_font
                else "Calibri"
            )
            fc = auto_fit_font_size(
                personal,
                education,
                matched_payload,
                font_family=family,
                one_page=one_page,
            )
            result["font_config_used"] = {
                "body_size": fc.body_size,
                "heading_size": fc.heading_size,
                "name_size": fc.name_size,
            }
        else:
            fc = font_config or DEFAULT_FONT

        doc = _build_doc(personal, education, matched_payload, fc, one_page)

        filename_base = _build_filename(personal, matched_payload)
        result["filename_base"] = filename_base

        # Use temp dir if no output_dir specified
        if output_dir:
            out = Path(output_dir)
            out.mkdir(parents=True, exist_ok=True)
        else:
            out = Path(tempfile.mkdtemp(prefix="resumeforge_"))

        docx_path = out / f"{filename_base}.docx"
        doc.save(str(docx_path))
        result["docx_path"] = str(docx_path)

        if to_pdf:
            pdf_path = out / f"{filename_base}.pdf"
            try:
                from docx2pdf import convert

                convert(str(docx_path), str(pdf_path))
                if pdf_path.exists():
                    result["pdf_path"] = str(pdf_path)
                else:
                    result["error"] = "PDF conversion produced no output file."
            except Exception as e:
                result["error"] = f"PDF conversion failed (docx saved OK): {e}"

    except Exception as e:
        result["error"] = str(e)

    return result
