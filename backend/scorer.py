"""
scorer.py
Scores a generated resume on two dimensions using Claude:

  1. ATS Score (0-10)   — how well the resume is formatted for ATS parsing
     (clean structure, keywords, no tables/graphics, proper sections, etc.)

  2. JD Match Score (0-10) — how well the resume content matches the job description
     (skills overlap, relevant experience, keyword density, role alignment)

Returns structured feedback so the UI can display score cards + actionable tips.
"""

import re
import json


def quick_gap_analysis(jd_text: str, resume_raw_text: str, client) -> dict:
    """
    Fast pre-generation analysis: compare JD requirements vs resume content.
    Returns keyword gaps the user can choose to include before generating.

    Returns:
    {
        "job_title": str,
        "company": str,
        "required_missing":  [{"keyword": str, "importance": "critical|high|medium", "explanation": str}],
        "preferred_missing": [{"keyword": str, "importance": "preferred", "explanation": str}],
        "already_have":      [str],   # skills/keywords already in resume
        "skill_categories":  {        # gap by category
            "Programming Languages": {"have": [...], "missing": [...]},
            ...
        },
        "summary": str,   # 1-sentence summary of gap
        "error": str|None
    }
    """
    prompt = f"""You are a resume gap analyst.

Compare this candidate's resume against the job description and identify skill/keyword gaps.

JOB DESCRIPTION:
{jd_text[:3500]}

CANDIDATE'S CURRENT RESUME:
{resume_raw_text[:6000]}

Return ONLY this JSON (no markdown, no explanation):
{{
  "job_title": "role title from JD",
  "company": "company from JD",
  "required_missing": [
    {{"keyword": "specific skill/tool", "importance": "critical", "explanation": "why this matters for the role"}}
  ],
  "preferred_missing": [
    {{"keyword": "specific skill/tool", "importance": "preferred", "explanation": "why this is a plus"}}
  ],
  "already_have": ["skill1", "skill2"],
  "skill_categories": {{
    "Category Name": {{
      "have": ["skill1", "skill2"],
      "missing": ["skill3"]
    }}
  }},
  "summary": "One sentence describing the overall gap"
}}

Rules:
- required_missing: skills explicitly marked as required/must-have in the JD but NOT in the resume
- preferred_missing: skills listed as preferred/nice-to-have but NOT in the resume
- already_have: top 10 JD-relevant skills the candidate clearly has
- skill_categories: 4-5 categories (Programming Languages, ML Frameworks, Cloud/MLOps, NLP/AI Tools, Other)
- Be specific — list actual tool/library names, not vague terms like "programming"
- Max 8 items each in required_missing and preferred_missing"""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = next((b.text for b in response.content if hasattr(b, "text")), "").strip()
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
        result = json.loads(raw)
        result["error"] = None
        return result
    except Exception as e:
        return {
            "job_title": "",
            "company": "",
            "required_missing": [],
            "preferred_missing": [],
            "already_have": [],
            "skill_categories": {},
            "summary": "",
            "error": str(e),
        }


def gap_analysis_markdown(gap: dict) -> str:
    """Format gap analysis as dark-mode-safe Markdown."""
    if gap.get("error"):
        return f"> ⚠ Gap analysis failed: {gap['error']}"

    lines = []

    have = gap.get("already_have", [])
    if have:
        lines += ["### ✅ Skills You Already Have (JD-relevant)", ""]
        lines.append(", ".join(f"**{s}**" for s in have[:12]))
        lines.append("")

    req = gap.get("required_missing", [])
    if req:
        lines += ["### ❌ Required Skills Missing from Your Resume", ""]
        lines += ["| Skill | Why It Matters |", "|---|---|"]
        for item in req:
            lines.append(f"| `{item['keyword']}` | {item.get('explanation', '')} |")
        lines.append("")

    pref = gap.get("preferred_missing", [])
    if pref:
        lines += ["### 🔶 Preferred Skills Not Yet Listed", ""]
        lines += ["| Skill | Why It Helps |", "|---|---|"]
        for item in pref:
            lines.append(f"| `{item['keyword']}` | {item.get('explanation', '')} |")
        lines.append("")

    cats = gap.get("skill_categories", {})
    if cats:
        lines += ["### 📊 Skills Breakdown by Category", ""]
        for cat, data in cats.items():
            have_list = ", ".join(data.get("have", [])) or "—"
            miss_list = ", ".join(data.get("missing", [])) or "—"
            lines.append(f"**{cat}**")
            lines.append(f"- Have: {have_list}")
            lines.append(f"- Missing: {miss_list}")
            lines.append("")

    return "\n".join(lines)


def _score_label(score: int) -> str:
    """Map a 0-10 integer score to a human-readable label."""
    if score >= 9:
        return "Excellent"
    if score >= 7:
        return "Strong"
    if score >= 5:
        return "Good"
    if score >= 3:
        return "Needs Work"
    return "Poor"


def _safe_list(value, default=None) -> list:
    """Return value as a list, falling back to default (empty list)."""
    if isinstance(value, list):
        return value
    return default if default is not None else []


def score_resume(resume_text: str, jd_text: str, client) -> dict:
    """
    Score the resume against the job description on two dimensions:
      - ATS Score  (0-10): structural quality, parsability, formatting
      - JD Match   (0-10): content alignment with this specific job description

    Args:
        resume_text: plain text extracted from the generated resume
        jd_text:     raw job description text
        client:      Anthropic client

    Returns:
        {
            "ats_score":        int (0-10),
            "ats_label":        str,
            "ats_feedback":     [str, str, str],   # actionable tips
            "match_score":      int (0-10),
            "match_label":      str,
            "match_feedback":   [str, str, str],
            "matched_keywords": [str, ...],
            "missing_keywords": [str, ...],
            "error":            str | None
        }
    """

    # ── Pre-compute keyword match rate (deterministic, not LLM-based) ──────────
    # Extract candidate keywords from the JD: words 4+ chars, skip stopwords.
    _STOPWORDS = {
        "with",
        "that",
        "this",
        "will",
        "have",
        "from",
        "they",
        "been",
        "their",
        "about",
        "more",
        "when",
        "your",
        "also",
        "into",
        "than",
        "then",
        "some",
        "what",
        "would",
        "which",
        "should",
        "could",
        "other",
        "work",
        "team",
        "you",
        "are",
        "for",
        "the",
        "and",
        "our",
        "not",
        "role",
        "must",
        "able",
        "experience",
        "including",
        "strong",
        "skills",
        "ability",
        "knowledge",
    }
    jd_words = re.findall(r"[a-zA-Z][a-zA-Z0-9+#.\-]{2,}", jd_text)
    jd_keywords = {
        w.lower() for w in jd_words if len(w) >= 4 and w.lower() not in _STOPWORDS
    }
    resume_lower = resume_text.lower()
    matched_count = sum(1 for kw in jd_keywords if kw in resume_lower)
    keyword_match_rate_pct = round(
        (matched_count / len(jd_keywords) * 100) if jd_keywords else 0, 1
    )

    prompt = f"""You are an expert ATS (Applicant Tracking System) analyst and resume coach.

Carefully evaluate the RESUME TEXT against the JOB DESCRIPTION and return honest, specific scores.

## JOB DESCRIPTION:
{jd_text[:5000]}

## RESUME TEXT:
{resume_text[:9000]}

## PRE-COMPUTED KEYWORD MATCH RATE
{keyword_match_rate_pct}% of unique JD keywords appear in the resume text.
Use this as a calibration anchor for both scores.

## SCORING DIMENSIONS

### 1. ATS Score (0-10)
How well is the resume structured for automated parsing by real ATS systems (Taleo, Workday, Greenhouse, Lever)?

Criteria:
- Standard section headings present: SUMMARY/PROFILE, SKILLS, EXPERIENCE, PROJECTS, EDUCATION (2 pts)
- Contact info complete and parseable: name, email, phone, LinkedIn/GitHub (1 pt)
- Consistent, unambiguous date formats (MM/YYYY or Month YYYY) throughout (1 pt)
- Bullet points start with strong action verbs (led, built, designed, improved, reduced…) (1 pt)
- Quantified achievements — numbers, percentages, impact metrics present (1 pt)
- Keyword match rate anchors this score: {keyword_match_rate_pct}% match rate
  - ≥70% → up to 4 pts, 50-69% → 3 pts, 30-49% → 2 pts, <30% → 0-1 pts (4 pts)

### 2. JD Match Score (0-10)
How well does the resume content match THIS specific job?

Criteria:
- Required technical skills from JD are present in resume (3 pts)
- Preferred/bonus skills from JD are covered (1 pt)
- Years of experience aligns with JD requirements (1 pt)
- Job title / seniority level is appropriate (1 pt)
- Project/work experience is relevant to the role domain (2 pts)
- Industry/domain terminology from JD appears in resume (1 pt)
- Overall narrative and positioning aligns with the role (1 pt)

### Calibration — be honest, not generous:
- 9-10 = Excellent (nearly perfect, very rare)
- 7-8 = Strong (clearly qualified, minor gaps)
- 5-6 = Good (qualified but meaningful gaps exist)
- 3-4 = Needs Work (significant gaps)
- 0-2 = Poor (major mismatch)

## OUTPUT FORMAT
Return ONLY valid JSON — no markdown fences, no prose before or after:
{{
  "ats_score": <integer 0-10>,
  "ats_label": "<Excellent|Strong|Good|Needs Work|Poor>",
  "ats_feedback": [
    "<specific, actionable improvement tip for ATS>",
    "<specific, actionable improvement tip for ATS>",
    "<specific, actionable improvement tip for ATS>"
  ],
  "match_score": <integer 0-10>,
  "match_label": "<Excellent|Strong|Good|Needs Work|Poor>",
  "match_feedback": [
    "<specific tip referencing an actual gap between resume and JD>",
    "<specific tip referencing an actual gap between resume and JD>",
    "<specific tip referencing an actual gap between resume and JD>"
  ],
  "matched_keywords": ["<JD keyword found in resume>", ...],
  "missing_keywords": ["<important JD keyword NOT in resume>", ...]
}}

Rules:
- matched_keywords: concrete skills/tools/technologies from the JD that appear in the resume (max 20)
- missing_keywords: important skills/tools from the JD that are absent from the resume (max 10)
- feedback tips must be specific — name the actual missing section, keyword, or skill
- do NOT return null for any array — use [] if empty
- scores must be integers, not strings or floats
- keyword_match_rate of {keyword_match_rate_pct}% should strongly inform ats_score"""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=1800,
            messages=[{"role": "user", "content": prompt}],
        )

        # Safely extract text from the first TextBlock in the response
        raw = ""
        for block in response.content:
            if hasattr(block, "text"):
                raw = block.text.strip()
                break

        if not raw:
            raise ValueError("Empty response from model")

        # Strip markdown fences if the model wraps in ```json ... ```
        raw = re.sub(r"^```[a-z]*\s*", "", raw)
        raw = re.sub(r"\s*```$", "", raw)
        raw = raw.strip()

        result = json.loads(raw)

        # Coerce and normalise all fields — Claude can return strings for ints, nulls for lists, etc.
        ats_score = max(0, min(10, int(result.get("ats_score", 0) or 0)))
        match_score = max(0, min(10, int(result.get("match_score", 0) or 0)))

        return {
            "ats_score": ats_score,
            "ats_label": _score_label(ats_score),
            "ats_feedback": _safe_list(result.get("ats_feedback")),
            "match_score": match_score,
            "match_label": _score_label(match_score),
            "match_feedback": _safe_list(result.get("match_feedback")),
            "matched_keywords": _safe_list(result.get("matched_keywords")),
            "missing_keywords": _safe_list(result.get("missing_keywords")),
            "error": None,
        }

    except Exception as e:
        return {
            "ats_score": 0,
            "ats_label": "N/A",
            "ats_feedback": [],
            "match_score": 0,
            "match_label": "N/A",
            "match_feedback": [],
            "matched_keywords": [],
            "missing_keywords": [],
            "error": str(e),
        }


def extract_resume_text(docx_path: str) -> str:
    """
    Extract all plain text from a generated .docx for scoring.
    Covers: body paragraphs, tables, headers, footers, and text boxes (shapes).
    """
    try:
        import docx
        from docx.oxml.ns import qn

        doc = docx.Document(docx_path)
        lines: list[str] = []

        def _para_text(para) -> str:
            return para.text.strip()

        # Body paragraphs
        for p in doc.paragraphs:
            t = _para_text(p)
            if t:
                lines.append(t)

        # Tables (resume builders sometimes use tables for layout)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for p in cell.paragraphs:
                        t = _para_text(p)
                        if t:
                            lines.append(t)

        # Headers and footers (contact info often lives here)
        for section in doc.sections:
            for hdr_ftr in (
                section.header,
                section.footer,
                section.even_page_header,
                section.even_page_footer,
                section.first_page_header,
                section.first_page_footer,
            ):
                try:
                    for p in hdr_ftr.paragraphs:
                        t = _para_text(p)
                        if t:
                            lines.append(t)
                except Exception:
                    pass

        # Text boxes / drawing shapes (some resume templates use these)
        body_xml = doc.element.body
        for txbx in body_xml.iter(qn("w:txbxContent")):
            for p_elem in txbx.iter(qn("w:p")):
                texts = [r.text for r in p_elem.iter(qn("w:t")) if r.text]
                combined = "".join(texts).strip()
                if combined:
                    lines.append(combined)

        return "\n".join(lines)
    except Exception:
        return ""


def score_card_markdown(scores: dict) -> str:
    """
    Render score results as Markdown — fully dark-mode compatible,
    no hardcoded background colours.
    """
    if scores.get("error") and not scores.get("ats_score"):
        return f"> ⚠ Scoring failed: {scores['error']}"

    def label_emoji(s):
        if s >= 9:
            return "🟢"
        if s >= 7:
            return "🟡"
        if s >= 5:
            return "🟠"
        return "🔴"

    ats = scores.get("ats_score", 0)
    match = scores.get("match_score", 0)
    ats_lbl = scores.get("ats_label", "")
    match_lbl = scores.get("match_label", "")

    lines = [
        "### Resume Scores",
        "",
        f"| | Score | Rating |",
        f"|---|---|---|",
        f"| **ATS Score** | {label_emoji(ats)} **{ats} / 10** | {ats_lbl} |",
        f"| **JD Match Score** | {label_emoji(match)} **{match} / 10** | {match_lbl} |",
        "",
    ]

    ats_tips = scores.get("ats_feedback", [])
    if ats_tips:
        lines.append("**ATS Tips**")
        for t in ats_tips[:3]:
            lines.append(f"- {t}")
        lines.append("")

    match_tips = scores.get("match_feedback", [])
    if match_tips:
        lines.append("**JD Match Tips**")
        for t in match_tips[:3]:
            lines.append(f"- {t}")
        lines.append("")

    matched = scores.get("matched_keywords", [])
    missing = scores.get("missing_keywords", [])
    if matched:
        lines.append(f"✅ **Keywords matched:** {', '.join(matched[:14])}")
    if missing:
        lines.append(f"❌ **Keywords missing:** {', '.join(missing[:10])}")

    return "\n".join(lines)


# Keep old name as alias so existing imports don't break
def score_card_html(scores: dict) -> str:
    return score_card_markdown(scores)

    ats_ring = ring(scores["ats_score"], scores["ats_label"], "ATS Score")
    match_ring = ring(scores["match_score"], scores["match_label"], "JD Match Score")

    def tips(items, icon="💡"):
        if not items:
            return ""
        lis = "".join(f"<li style='margin:4px 0'>{icon} {t}</li>" for t in items[:3])
        return f"<ul style='margin:6px 0 0;padding-left:1.2rem;font-size:0.88rem;color:#444'>{lis}</ul>"

    matched = ", ".join(scores.get("matched_keywords", [])[:12])
    missing = ", ".join(scores.get("missing_keywords", [])[:8])

    kw_section = ""
    if matched or missing:
        kw_section = f"""
        <div style="margin-top:12px;font-size:0.85rem">
          {"<div style='margin-bottom:4px'><span style='color:#1A7A3F;font-weight:700'>✓ Keywords matched:</span> " + matched + "</div>" if matched else ""}
          {"<div><span style='color:#C0392B;font-weight:700'>✗ Keywords missing:</span> " + missing + "</div>" if missing else ""}
        </div>"""

    return f"""
    <div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:1.2rem 1.4rem;margin-top:1rem">
      <div style="font-size:1rem;font-weight:700;color:#1A2E4A;margin-bottom:1rem">Resume Scores</div>
      <div style="display:flex;gap:2rem;flex-wrap:wrap;justify-content:flex-start;align-items:flex-start">
        {ats_ring}
        {match_ring}
        <div style="flex:3;min-width:220px">
          <div style="font-size:0.85rem;font-weight:700;color:#555;margin-bottom:2px">ATS Tips</div>
          {tips(scores.get("ats_feedback", []), "✦")}
          <div style="font-size:0.85rem;font-weight:700;color:#555;margin-top:10px;margin-bottom:2px">Match Tips</div>
          {tips(scores.get("match_feedback", []), "✦")}
          {kw_section}
        </div>
      </div>
    </div>"""
