"""
scorer.py
Scores a generated resume on two dimensions using Claude:

  1. ATS Score (0-10)   — how well the resume is formatted for ATS parsing
     (clean structure, keywords, no tables/graphics, proper sections, etc.)

  2. JD Match Score (0-10) — how well the resume content matches the job description
     (skills overlap, relevant experience, keyword density, role alignment)

Returns structured feedback so the UI can display score cards + actionable tips.
"""

import re
import json


def quick_gap_analysis(jd_text: str, resume_raw_text: str, client) -> dict:
    """
    Fast pre-generation analysis: compare JD requirements vs resume content.
    Returns keyword gaps the user can choose to include before generating.

    Returns:
    {
        "job_title": str,
        "company": str,
        "required_missing":  [{"keyword": str, "importance": "critical|high|medium", "explanation": str}],
        "preferred_missing": [{"keyword": str, "importance": "preferred", "explanation": str}],
        "already_have":      [str],   # skills/keywords already in resume
        "skill_categories":  {        # gap by category
            "Programming Languages": {"have": [...], "missing": [...]},
            ...
        },
        "summary": str,   # 1-sentence summary of gap
        "error": str|None
    }
    """
    prompt = f"""You are a resume gap analyst.

Compare this candidate's resume against the job description and identify skill/keyword gaps.

JOB DESCRIPTION:
{jd_text[:3500]}

CANDIDATE'S CURRENT RESUME:
{resume_raw_text[:6000]}

Return ONLY this JSON (no markdown, no explanation):
{{
  "job_title": "role title from JD",
  "company": "company from JD",
  "required_missing": [
    {{"keyword": "specific skill/tool", "importance": "critical", "explanation": "why this matters for the role"}}
  ],
  "preferred_missing": [
    {{"keyword": "specific skill/tool", "importance": "preferred", "explanation": "why this is a plus"}}
  ],
  "already_have": ["skill1", "skill2"],
  "skill_categories": {{
    "Category Name": {{
      "have": ["skill1", "skill2"],
      "missing": ["skill3"]
    }}
  }},
  "summary": "One sentence describing the overall gap"
}}

Rules:
- required_missing: skills explicitly marked as required/must-have in the JD but NOT in the resume
- preferred_missing: skills listed as preferred/nice-to-have but NOT in the resume
- already_have: top 10 JD-relevant skills the candidate clearly has
- skill_categories: 4-5 categories (Programming Languages, ML Frameworks, Cloud/MLOps, NLP/AI Tools, Other)
- Be specific — list actual tool/library names, not vague terms like "programming"
- Max 8 items each in required_missing and preferred_missing"""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=1500,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
        result = json.loads(raw)
        result["error"] = None
        return result
    except Exception as e:
        return {
            "job_title": "",
            "company": "",
            "required_missing": [],
            "preferred_missing": [],
            "already_have": [],
            "skill_categories": {},
            "summary": "",
            "error": str(e),
        }


def gap_analysis_markdown(gap: dict) -> str:
    """Format gap analysis as dark-mode-safe Markdown."""
    if gap.get("error"):
        return f"> ⚠ Gap analysis failed: {gap['error']}"

    lines = []

    have = gap.get("already_have", [])
    if have:
        lines += ["### ✅ Skills You Already Have (JD-relevant)", ""]
        lines.append(", ".join(f"**{s}**" for s in have[:12]))
        lines.append("")

    req = gap.get("required_missing", [])
    if req:
        lines += ["### ❌ Required Skills Missing from Your Resume", ""]
        lines += ["| Skill | Why It Matters |", "|---|---|"]
        for item in req:
            lines.append(f"| `{item['keyword']}` | {item.get('explanation', '')} |")
        lines.append("")

    pref = gap.get("preferred_missing", [])
    if pref:
        lines += ["### 🔶 Preferred Skills Not Yet Listed", ""]
        lines += ["| Skill | Why It Helps |", "|---|---|"]
        for item in pref:
            lines.append(f"| `{item['keyword']}` | {item.get('explanation', '')} |")
        lines.append("")

    cats = gap.get("skill_categories", {})
    if cats:
        lines += ["### 📊 Skills Breakdown by Category", ""]
        for cat, data in cats.items():
            have_list = ", ".join(data.get("have", [])) or "—"
            miss_list = ", ".join(data.get("missing", [])) or "—"
            lines.append(f"**{cat}**")
            lines.append(f"- Have: {have_list}")
            lines.append(f"- Missing: {miss_list}")
            lines.append("")

    return "\n".join(lines)


def score_resume(resume_text: str, jd_text: str, client) -> dict:
    """
    Score the resume against the job description.

    Args:
        resume_text: plain text extracted from the generated resume
        jd_text:     raw job description text
        client:      Anthropic client

    Returns:
        {
            "ats_score":        int (0-10),
            "ats_label":        str (e.g. "Strong"),
            "ats_feedback":     [list of 3 short tips],
            "match_score":      int (0-10),
            "match_label":      str (e.g. "Good Match"),
            "match_feedback":   [list of 3 short tips],
            "matched_keywords": [list of JD keywords found in resume],
            "missing_keywords": [list of important JD keywords not in resume],
            "error":            str | None
        }
    """

    prompt = f"""You are an expert ATS (Applicant Tracking System) analyst and resume coach.

Evaluate this resume against the job description and return scores + feedback.

## JOB DESCRIPTION:
{jd_text[:4000]}

## RESUME TEXT:
{resume_text[:8000]}

## YOUR TASK:
Score the resume on two dimensions, both out of 10:

### ATS Score (0-10)
Measures how well the resume is formatted for ATS parsing:
- Proper section headings (SKILLS, EXPERIENCE, PROJECTS, EDUCATION)
- No graphics, tables, columns, or text boxes that confuse parsers
- Action verbs, quantified achievements
- Appropriate keyword density (not stuffed)
- Contact info clearly present
- Standard date formats

### JD Match Score (0-10)
Measures how well the resume content aligns with this specific JD:
- Required skills coverage
- Preferred skills coverage
- Relevant experience and project overlap
- Keyword and terminology match
- Role-level alignment

### Score labels:
- 9-10: Excellent
- 7-8: Strong  
- 5-6: Good
- 3-4: Needs Work
- 0-2: Poor

Return ONLY this JSON (no markdown fences, no explanation):
{{
  "ats_score": <int 0-10>,
  "ats_label": "<label>",
  "ats_feedback": [
    "<specific actionable tip 1>",
    "<specific actionable tip 2>",
    "<specific actionable tip 3>"
  ],
  "match_score": <int 0-10>,
  "match_label": "<label>",
  "match_feedback": [
    "<specific actionable tip 1>",
    "<specific actionable tip 2>",
    "<specific actionable tip 3>"
  ],
  "matched_keywords": ["<keyword1>", "<keyword2>", ...],
  "missing_keywords": ["<keyword1>", "<keyword2>", ...]
}}"""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=1200,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
        result = json.loads(raw)
        result["error"] = None
        return result
    except Exception as e:
        return {
            "ats_score": 0,
            "ats_label": "N/A",
            "ats_feedback": [],
            "match_score": 0,
            "match_label": "N/A",
            "match_feedback": [],
            "matched_keywords": [],
            "missing_keywords": [],
            "error": str(e),
        }


def extract_resume_text(docx_path: str) -> str:
    """Extract plain text from the generated .docx for scoring."""
    try:
        import docx

        doc = docx.Document(docx_path)
        lines = [p.text for p in doc.paragraphs if p.text.strip()]
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        lines.append(cell.text.strip())
        return "\n".join(lines)
    except Exception:
        return ""


def score_card_markdown(scores: dict) -> str:
    """
    Render score results as Markdown — fully dark-mode compatible,
    no hardcoded background colours.
    """
    if scores.get("error") and not scores.get("ats_score"):
        return f"> ⚠ Scoring failed: {scores['error']}"

    def label_emoji(s):
        if s >= 9:
            return "🟢"
        if s >= 7:
            return "🟡"
        if s >= 5:
            return "🟠"
        return "🔴"

    ats = scores.get("ats_score", 0)
    match = scores.get("match_score", 0)
    ats_lbl = scores.get("ats_label", "")
    match_lbl = scores.get("match_label", "")

    lines = [
        "### Resume Scores",
        "",
        f"| | Score | Rating |",
        f"|---|---|---|",
        f"| **ATS Score** | {label_emoji(ats)} **{ats} / 10** | {ats_lbl} |",
        f"| **JD Match Score** | {label_emoji(match)} **{match} / 10** | {match_lbl} |",
        "",
    ]

    ats_tips = scores.get("ats_feedback", [])
    if ats_tips:
        lines.append("**ATS Tips**")
        for t in ats_tips[:3]:
            lines.append(f"- {t}")
        lines.append("")

    match_tips = scores.get("match_feedback", [])
    if match_tips:
        lines.append("**JD Match Tips**")
        for t in match_tips[:3]:
            lines.append(f"- {t}")
        lines.append("")

    matched = scores.get("matched_keywords", [])
    missing = scores.get("missing_keywords", [])
    if matched:
        lines.append(f"✅ **Keywords matched:** {', '.join(matched[:14])}")
    if missing:
        lines.append(f"❌ **Keywords missing:** {', '.join(missing[:10])}")

    return "\n".join(lines)


# Keep old name as alias so existing imports don't break
def score_card_html(scores: dict) -> str:
    return score_card_markdown(scores)

    ats_ring = ring(scores["ats_score"], scores["ats_label"], "ATS Score")
    match_ring = ring(scores["match_score"], scores["match_label"], "JD Match Score")

    def tips(items, icon="💡"):
        if not items:
            return ""
        lis = "".join(f"<li style='margin:4px 0'>{icon} {t}</li>" for t in items[:3])
        return f"<ul style='margin:6px 0 0;padding-left:1.2rem;font-size:0.88rem;color:#444'>{lis}</ul>"

    matched = ", ".join(scores.get("matched_keywords", [])[:12])
    missing = ", ".join(scores.get("missing_keywords", [])[:8])

    kw_section = ""
    if matched or missing:
        kw_section = f"""
        <div style="margin-top:12px;font-size:0.85rem">
          {"<div style='margin-bottom:4px'><span style='color:#1A7A3F;font-weight:700'>✓ Keywords matched:</span> " + matched + "</div>" if matched else ""}
          {"<div><span style='color:#C0392B;font-weight:700'>✗ Keywords missing:</span> " + missing + "</div>" if missing else ""}
        </div>"""

    return f"""
    <div style="background:#fff;border:1px solid #e0e0e0;border-radius:10px;padding:1.2rem 1.4rem;margin-top:1rem">
      <div style="font-size:1rem;font-weight:700;color:#1A2E4A;margin-bottom:1rem">Resume Scores</div>
      <div style="display:flex;gap:2rem;flex-wrap:wrap;justify-content:flex-start;align-items:flex-start">
        {ats_ring}
        {match_ring}
        <div style="flex:3;min-width:220px">
          <div style="font-size:0.85rem;font-weight:700;color:#555;margin-bottom:2px">ATS Tips</div>
          {tips(scores.get("ats_feedback", []), "✦")}
          <div style="font-size:0.85rem;font-weight:700;color:#555;margin-top:10px;margin-bottom:2px">Match Tips</div>
          {tips(scores.get("match_feedback", []), "✦")}
          {kw_section}
        </div>
      </div>
    </div>"""
