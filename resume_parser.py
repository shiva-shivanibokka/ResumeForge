"""
resume_parser.py
Extracts structured data from an uploaded resume (PDF or .docx).

Returns a dict with:
  name, email, phone, linkedin, github, location,
  education (list), experience (list), skills (list), raw_text
"""

import re
import io
from pathlib import Path


# ──────────────────────────────────────────────────────────────────────────────
# TEXT EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────


def extract_text_from_pdf(file_bytes: bytes) -> str:
    """Extract plain text from a PDF using PyMuPDF (fast, layout-aware)."""
    try:
        import fitz  # PyMuPDF

        doc = fitz.open(stream=file_bytes, filetype="pdf")
        pages = []
        for page in doc:
            pages.append(page.get_text("text"))
        return "\n".join(pages)
    except Exception as e:
        # Fallback to pdfminer
        try:
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams

            output = io.StringIO()
            extract_text_to_fp(
                io.BytesIO(file_bytes),
                output,
                laparams=LAParams(),
                output_type="text",
                codec="utf-8",
            )
            return output.getvalue()
        except Exception as e2:
            return f"[PDF extraction failed: {e} / {e2}]"


def extract_text_from_docx(file_bytes: bytes) -> str:
    """Extract plain text from a .docx file."""
    try:
        import docx

        doc = docx.Document(io.BytesIO(file_bytes))
        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # Also grab text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():
                        paragraphs.append(cell.text.strip())
        return "\n".join(paragraphs)
    except Exception as e:
        return f"[DOCX extraction failed: {e}]"


def extract_raw_text(file_path: str, file_bytes: bytes = None) -> str:
    """
    Extract raw text from a resume file.
    Accepts either a file path string or raw bytes + extension hint.
    """
    path = Path(file_path)
    if file_bytes is None:
        file_bytes = path.read_bytes()

    ext = path.suffix.lower()
    if ext == ".pdf":
        return extract_text_from_pdf(file_bytes)
    elif ext in (".docx", ".doc"):
        return extract_text_from_docx(file_bytes)
    else:
        # Try to decode as plain text
        try:
            return file_bytes.decode("utf-8", errors="ignore")
        except Exception:
            return ""


# ──────────────────────────────────────────────────────────────────────────────
# REGEX FIELD EXTRACTION  (fast, no API call needed)
# ──────────────────────────────────────────────────────────────────────────────


def _find_email(text: str) -> str:
    match = re.search(r"[a-zA-Z0-9._%+\-]+@[a-zA-Z0-9.\-]+\.[a-zA-Z]{2,}", text)
    return match.group(0) if match else ""


def _find_phone(text: str) -> str:
    match = re.search(r"(\+?1[\s\-.]?)?\(?\d{3}\)?[\s\-.]?\d{3}[\s\-.]?\d{4}", text)
    return match.group(0).strip() if match else ""


def _find_linkedin(text: str) -> str:
    match = re.search(r"linkedin\.com/in/[\w\-]+", text, re.IGNORECASE)
    return match.group(0) if match else ""


def _find_github(text: str) -> str:
    match = re.search(r"github\.com/[\w\-]+", text, re.IGNORECASE)
    return match.group(0) if match else ""


def _find_name(text: str) -> str:
    """
    Heuristic: the name is usually on the first non-empty line,
    often as 2-4 capitalized words.
    """
    for line in text.splitlines()[:10]:
        line = line.strip()
        if not line:
            continue
        # Skip lines that look like contact info
        if any(c in line for c in ["@", "http", "linkedin", "github", "+1", "("]):
            continue
        # Match 2-4 capitalized words (name pattern)
        if re.match(r"^([A-Z][a-z]+\s){1,3}[A-Z][a-z]+$", line):
            return line
        # Also accept ALL-CAPS names
        if re.match(r"^[A-Z\s]{4,40}$", line) and len(line.split()) >= 2:
            return line.title()
    return ""


# ──────────────────────────────────────────────────────────────────────────────
# CLAUDE-POWERED STRUCTURED EXTRACTION
# ──────────────────────────────────────────────────────────────────────────────


def parse_resume_with_claude(raw_text: str, client) -> dict:
    """
    Use Claude to extract fully structured data from resume text.
    Falls back to regex for contact info if Claude fails.
    """
    import json

    prompt = f"""You are a resume parsing assistant. Extract structured information from this resume text.

Return a JSON object with exactly these keys:
{{
  "name": "Full name",
  "email": "email address",
  "phone": "phone number",
  "linkedin": "LinkedIn URL or username",
  "github": "GitHub URL or username",
  "location": "City, State or City, Country",
  "education": [
    {{
      "school": "Institution name",
      "degree": "Degree and field of study",
      "dates": "Start – End",
      "gpa": "GPA if listed, else null",
      "location": "City, State"
    }}
  ],
  "experience": [
    {{
      "company": "Company name",
      "title": "Job title",
      "dates": "Start – End",
      "location": "City, State",
      "bullets": ["bullet 1", "bullet 2", ...]
    }}
  ],
  "skills": ["skill1", "skill2", ...],
  "summary": "Professional summary if present, else empty string"
}}

Only return the JSON. No markdown fences. No explanation.

RESUME TEXT:
{raw_text[:8000]}"""

    try:
        response = client.messages.create(
            model="claude-opus-4-5",
            max_tokens=3000,
            messages=[{"role": "user", "content": prompt}],
        )
        raw = response.content[0].text.strip()
        raw = re.sub(r"^```[a-z]*\n?", "", raw)
        raw = re.sub(r"\n?```$", "", raw)
        parsed = json.loads(raw)
    except Exception:
        # Fallback: build minimal structure from regex
        parsed = {
            "name": _find_name(raw_text),
            "email": _find_email(raw_text),
            "phone": _find_phone(raw_text),
            "linkedin": _find_linkedin(raw_text),
            "github": _find_github(raw_text),
            "location": "",
            "education": [],
            "experience": [],
            "skills": [],
            "summary": "",
        }

    # Always patch contact fields with regex if Claude missed them
    if not parsed.get("email"):
        parsed["email"] = _find_email(raw_text)
    if not parsed.get("phone"):
        parsed["phone"] = _find_phone(raw_text)
    if not parsed.get("linkedin"):
        parsed["linkedin"] = _find_linkedin(raw_text)
    if not parsed.get("github"):
        parsed["github"] = _find_github(raw_text)
    if not parsed.get("name"):
        parsed["name"] = _find_name(raw_text)

    parsed["raw_text"] = raw_text
    return parsed


def parse_resume(file_path: str, client, file_bytes: bytes = None) -> dict:
    """
    Full pipeline: extract text → parse with Claude → return structured dict.
    """
    raw_text = extract_raw_text(file_path, file_bytes)
    if not raw_text or raw_text.startswith("["):
        return {
            "name": "",
            "email": "",
            "phone": "",
            "linkedin": "",
            "github": "",
            "location": "",
            "education": [],
            "experience": [],
            "skills": [],
            "summary": "",
            "raw_text": raw_text,
            "error": raw_text,
        }
    return parse_resume_with_claude(raw_text, client)
